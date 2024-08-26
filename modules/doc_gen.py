from docx import Document
from logging import debug
from os import makedirs


class Novel:
    def __init__(self, novel: dict):
        self.title = self._set_title(novel['urls'][0])
        self.length = len(novel['titles'])

        self.contents = list(zip(novel['titles'], novel['chr_contents']))
        self.doc_name = None

    def _set_title(self, title: str):
        title = title.split("/")[-2]
        title = title.replace('-', ' ').title()
        return title

    def _set_doc_name(self, doc_lang: str):
        words = self.title.split()
        # Extrair a primeira letra de cada palavra e transformar em maiúscula
        doc_name = ''.join(word[0] for word in words)
        self.doc_name = doc_lang + '_' + doc_name


def run_docx_generate(novel_info: dict, docs_number: int, doc_lang: str):
    novel = Novel(novel_info)
    novel._set_doc_name(doc_lang)

    debug('Verificando se há diretório para a novel...')
    makedirs(f'Documentos/{novel.title}', exist_ok=True)

    debug(f'Nome da novel: {novel.title}')
    debug(f'Número de capítulos: {novel.length}')
    debug(f'Número de documentos a serem gerados: {docs_number}')

    # Pelo menos um capítulo por documento.
    chapters_per_doc = max(novel.length // docs_number, 1)
    debug(f'Número de capítulos por documento: {chapters_per_doc}')

    generate_docs(docs_number, novel, chapters_per_doc)

    debug('Documentos gerados com sucesso.')


def generate_docs(docs_number: int, novel: Novel, chapters_per_doc: int):
    for i in range(docs_number):
        start = i * chapters_per_doc
        end = start + chapters_per_doc

        debug(f'Start: {start}, End: {end - 1}')

        doc = Document()
        if i == docs_number - 1:
            end = novel.length

        doc_info = add_chapters(doc, novel, start, end)
        doc, first, last = doc_info
        doc.save(
            f'Documentos/{novel.title}/{novel.doc_name}_{first}-{last}.docx')

        debug(f'Documento {i + 1} gerado.')


def add_chapters(doc: Document, novel: Novel, start: int, end: int):
    first, last = get_first_last(novel, start, end)

    count = 0
    for i, (chr_title, chr_content) in enumerate(novel.contents[start:end]):
        doc.add_heading(chr_title, level=1)
        for paragraph in chr_content.split('\n'):
            doc.add_paragraph(paragraph)
        doc.add_page_break()

        # Notifica a cada 50 capítulos adicionados.
        if i != 0 and ((i + 1) % 50 == 0 or i == end - 1):
            count += 1
            num = count * 50
            if i == end - 1:
                num = i + 1
            debug(f'Adicionados {num} capítulos.')
    return doc, first, last


def get_first_last(novel: Novel, start: int, end: int):
    first_chr = novel.contents[start][0]
    last_chr = novel.contents[end - 1][0]
    debug(f'Adicionando capítulos de "{first_chr}" até "{last_chr}".')

    # Se estiver em português, traduz para inglês.
    first_chr = first_chr.replace('Capítulo', 'Chapter')
    last_chr = last_chr.replace('Capítulo', 'Chapter')

    if 'chapter' in first_chr.lower() or 'chapter' in last_chr.lower():
        if 'chapter' in first_chr.lower():
            first_chr = first_chr.replace(':', ' ').split(' ')[1]
            first_chr = int(first_chr)
        if 'chapter' in last_chr.lower():
            last_chr = last_chr.replace(':', ' ').split(' ')[1]
            last_chr = int(last_chr)
    else:
        first_chr = first_chr.lower()
        last_chr = last_chr.lower()

    return first_chr, last_chr
